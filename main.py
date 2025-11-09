import PyPDF2
import docx
import io
from typing import List, Dict, Optional, Union, BinaryIO
import logging
from pathlib import Path
import tempfile
import os
from parser import ResumeParser
from screener import ResumeScreener
from resume_optimizer import ResumeOptimizer
import utils

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# ...existing code...
class ResumeExtractor:
    """Handles extraction of text and URLs from resume files (PDF, DOCX)."""

    @staticmethod
    def extract_urls_from_pdf(file_path: str) -> List[str]:
        """Extract URLs from PDF hyperlinks/annotations."""
        urls = []
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    annotations = page.get("/Annots")
                    if not annotations:
                        continue

                    # Ensure iterable
                    if not isinstance(annotations, list):
                        annotations = [annotations]

                    for annotation_ref in annotations:
                        try:
                            annotation = annotation_ref.get_object()
                        except Exception:
                            continue

                        if annotation.get("/Subtype") != "/Link":
                            continue

                        action = annotation.get("/A")
                        if not action:
                            continue

                        try:
                            action = action.get_object()
                        except Exception:
                            pass

                        uri = action.get("/URI") if isinstance(action, dict) else None
                        if uri:
                            url = uri if isinstance(uri, str) else str(uri)
                            if url not in urls:
                                urls.append(url)
        except Exception as e:
            logger.warning(f"Could not extract URLs from PDF: {e}")

        return urls
    
    @staticmethod
    def extract_urls_from_docx(file_path: str) -> List[str]:
        """Extract URLs from DOCX hyperlinks."""
        urls = []
        try:
            doc = docx.Document(file_path)
            # Get all hyperlinks from the document relationships
            rels = doc.part.rels
            for rel in rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel.target_ref
                    if url and url not in urls:
                        # Filter out internal anchors (starting with #)
                        if not url.startswith("#"):
                            urls.append(url)
        except Exception as e:
            logger.warning(f"Could not extract URLs from DOCX: {e}")

        return urls
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text and URLs from PDF file."""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text_parts: List[str] = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                except Exception:
                    page_text = None
                text_parts.append(page_text or "")

            text = "\n".join(text_parts)

            # Also extract URLs from hyperlinks
            urls = ResumeExtractor.extract_urls_from_pdf(file_path)
            if urls:
                text += "\n\nEXTRACTED URLS/LINKS:\n"
                for url in urls:
                    text += f"- {url}\n"

            return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text and URLs from DOCX file."""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Also extract URLs from hyperlinks
        urls = ResumeExtractor.extract_urls_from_docx(file_path)
        if urls:
            text += "\n\nEXTRACTED URLS/LINKS:\n"
            for url in urls:
                text += f"- {url}\n"

        return text

    @staticmethod
    def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
        """Extract text and URLs from PDF bytes (for file uploads)."""
        # Write bytes to a temporary file so URL extraction (which relies on file path)
        # can reuse file-based logic that resolves annotations reliably.
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            # Use the file-based extractor (it calls extract_urls_from_pdf)
            return ResumeExtractor.extract_text_from_pdf(tmp_path)
        finally:
            if tmp_path:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    @staticmethod
    def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
        """Extract text and URLs from DOCX bytes (for file uploads)."""
        file_obj = io.BytesIO(file_bytes)
        doc = docx.Document(file_obj)
        text = "\n".join([paragraph.text or "" for paragraph in doc.paragraphs])
        return text
    # ...existing code...
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        """
        Extract text from file bytes (for file uploads).
        Auto-detects file type based on filename extension.
        """
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            logger.info(f"Extracting text from uploaded PDF: {filename}")
            return ResumeExtractor.extract_text_from_pdf_bytes(file_bytes)
        elif extension in [".docx", ".doc"]:
            logger.info(f"Extracting text from uploaded DOCX: {filename}")
            return ResumeExtractor.extract_text_from_docx_bytes(file_bytes)
        else:
            raise ValueError(
                f"Unsupported file format: {extension}. Supported formats: .pdf, .docx, .doc"
            )

    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        """
        Extract text from a resume file (PDF or DOCX).
        Auto-detects file type based on extension.
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path_obj.suffix.lower()

        if extension == ".pdf":
            logger.info(f"Extracting text from PDF: {file_path}")
            return ResumeExtractor.extract_text_from_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            logger.info(f"Extracting text from DOCX: {file_path}")
            return ResumeExtractor.extract_text_from_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {extension}. Supported formats: .pdf, .docx, .doc"
            )

class ResumeProcessor:
    """
    Main class for processing resumes - parsing and screening.
    Designed for easy integration with Flask/FastAPI backends.
    """
    @staticmethod
    def _normalize_screening_result(result: Dict) -> Dict:
        """
        Ensure screening result fields expected to be objects are objects.
        Some downstream validation expects structured objects for match
        fields; if the screener produced plain strings (or lists), wrap
        them into small objects to satisfy the schema and keep the text.
        """
        if not isinstance(result, dict):
            return result

        # Keys that validation expects to be objects (not raw strings)
        object_keys = [
            "project_match",
            "education_match",
            "experience_match",
            "skill_match",
            "cultural_fit",
        ]

        for key in object_keys:
            if key in result:
                val = result[key]
                if isinstance(val, str):
                    result[key] = {"text": val}
                elif isinstance(val, list):
                    # keep original items but embed them
                    result[key] = {"items": val}
                elif val is None:
                    result[key] = {}
                # if already a dict/object, leave as-is
        # Ensure overall_score is numeric (some tools expect a number)
        if "overall_score" in result:
            try:
                # preserve numeric if already numeric; coerce strings that look like numbers
                if isinstance(result["overall_score"], str):
                    result["overall_score"] = float(result["overall_score"])
            except Exception:
                # if coercion fails, drop it to avoid validation issues
                result.pop("overall_score", None)

        return result
    
    def __init__(
        self,
        parser_model: str = "llama-3.3-70b-versatile",
        screener_model: str = "llama-3.3-70b-versatile",
        optimizer_model: str = "llama-3.3-70b-versatile",
    ):
        """
        Initialize the resume processor.

        Args:
            parser_model: Groq model for parsing resumes
            screener_model: Groq model for screening resumes
        """
        self.parser = ResumeParser(model=parser_model)
        self.screener = ResumeScreener(model=screener_model)
        self.optimizer = ResumeOptimizer(model=optimizer_model)
        self.extractor = ResumeExtractor()

    def parse_resume_from_path(self, file_path: str) -> Dict:
        """
        Parse resume from file path.

        Args:
            file_path: Path to resume file

        Returns:
            Parsed resume as dictionary
        """
        resume_text = self.extractor.extract_text_from_file(file_path)
        logger.info(f"Extracted {len(resume_text)} characters from {file_path}")

        parsed_resume = self.parser.parse_resume(resume_text)
        logger.info("Resume parsed successfully")

        return parsed_resume.model_dump(exclude_none=True)

    def parse_resume_from_bytes(self, file_bytes: bytes, filename: str) -> Dict:
        """
        Parse resume from file bytes (for file uploads).

        Args:
            file_bytes: File content as bytes
            filename: Original filename (for extension detection)

        Returns:
            Parsed resume as dictionary
        """
        resume_text = self.extractor.extract_text_from_bytes(file_bytes, filename)
        logger.info(
            f"Extracted {len(resume_text)} characters from uploaded file: {filename}"
        )

        parsed_resume = self.parser.parse_resume(resume_text)
        logger.info("Resume parsed successfully")

        return parsed_resume.model_dump(exclude_none=True)

    def screen_resume(
        self,
        parsed_resume: Dict,
        job_title: str,
        job_description: str,
        weights: Optional[Dict[str, float]] = None,
    ) -> Dict:
        """
        Screen a parsed resume against job requirements.

        Args:
            parsed_resume: Parsed resume dictionary
            job_title: Job position title
            job_description: Job description and requirements
            weights: Optional custom weights for scoring categories

        Returns:
            Screening result as dictionary
        """
        # Normalize resume skills
        parsed_resume["skills"] = utils.normalize_skills(parsed_resume.get("skills", []))
    
        # Expand to include semantically similar ones found in JD
        parsed_resume["skills"] = utils.fuzzy_expand_skills(parsed_resume["skills"], job_description)
    
        screening_result = self.screener.screen_resume(
            parsed_resume, job_title, job_description, weights
        )
        logger.info(
            f"Resume screened. Overall score: {screening_result.overall_score}/10"
        )
        # Accept either an object with model_dump or a plain dict from the screener
        if hasattr(screening_result, "model_dump"):
            result_dict = screening_result.model_dump(exclude_none=True)
        elif isinstance(screening_result, dict):
            result_dict = screening_result
        else:
            # fallback: try to convert to dict safely
            try:
                result_dict = dict(screening_result)
            except Exception:
                result_dict = {"error": "invalid screening result format"}

        # Normalize fields that validation/tools expect to be objects
        result_dict = self._normalize_screening_result(result_dict)

        # Log overall score when present and numeric
        overall = result_dict.get("overall_score")
        try:
            logger.info(f"Resume screened. Overall score: {overall}/10")
        except Exception:
            logger.info("Resume screened.")

        return result_dict

        # return screening_result.model_dump(exclude_none=True)

    def optimise_resume(self, parsed_resume: Dict, job_title: str, job_description: str, screening_result: Dict) -> Dict:
        """
        Optimise the resume by providing suggestions.

        Args:
            parsed_resume: Parsed resume dictionary
            job_title: Job position title
            job_description: Job description and requirements
            screening_result: Result from screening the resume

        Returns:
            Optimization suggestions as dictionary
        """
        optimization = self.optimizer.generate_suggestions(parsed_resume, job_title, job_description, screening_result)
        return optimization.model_dump(exclude_none=True)

    def process_resume_from_path(
        self,
        file_path: str,
        job_title: str,
        job_description: str,
        weights: Optional[Dict[str, float]] = None,
    ) -> Dict:
        """
        Complete workflow: Parse and screen resume from file path.

        Args:
            file_path: Path to resume file
            job_title: Job position title
            job_description: Job description and requirements
            weights: Optional custom weights for scoring

        Returns:
            Dictionary with 'parsed' and 'screened' keys
        """
        logger.info(f"Processing resume from path: {file_path}")

        # Parse resume
        parsed = self.parse_resume_from_path(file_path)

        # Screen resume
        screened = self.screen_resume(parsed, job_title, job_description, weights)

        # Optimise the resume by providing suggestions
        optimization = self.optimise_resume(parsed, job_title, job_description, screened)

        return {"parsed": parsed, "screened": screened, "optimization": optimization}

    def process_resume_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        job_title: str,
        job_description: str,
        weights: Optional[Dict[str, float]] = None,
    ) -> Dict:
        """
        Complete workflow: Parse and screen resume from file bytes (for uploads).
        Perfect for Flask/FastAPI endpoints.

        Args:
            file_bytes: File content as bytes
            filename: Original filename
            job_title: Job position title
            job_description: Job description and requirements
            weights: Optional custom weights for scoring

        Returns:
            Dictionary with 'parsed' and 'screened' keys
        """
        logger.info(f"Processing uploaded resume: {filename}")

        # Parse resume
        parsed = self.parse_resume_from_bytes(file_bytes, filename)

        # Screen resume
        screened = self.screen_resume(parsed, job_title, job_description, weights)

       # Optimise the resume by providing suggestions
        optimization = self.optimizer.generate_suggestions(parsed, job_title, job_description, screened)

        return {"parsed": parsed, "screened": screened, "optimization": optimization}